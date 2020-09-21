
from docx import Document
from docx.shared import Inches
import os

def doc_template():
    doc = Document()
    doc.add_heading('Il mio primo report ', 0)
    folder = "static/report/"

    todisplaylist = ["andamento_nazionale.png","variazione_totale_positivi.png"]
    for name in todisplaylist:
        path = os.path.join(folder, name)
        doc.add_picture(path, width=Inches(6))

    doc.save("text.docx")

if __name__ == "__main__":
    doc_template()