import numpy as np
#from flask_ngrok import run_with_ngrok
from flask import Flask,render_template,request,send_file
from utils import *
import io
from docx import Document
from docx.shared import Inches
import os


app = Flask(__name__)
#run_with_ngrok(app)  # Start ngrok when app is run

def load_data():
    df = pd.read_csv("https://raw.githubusercontent.com/pcm-dpc/COVID-19/master/dati-andamento-nazionale/dpc-covid19-ita-andamento-nazionale.csv")
    df["data"] = [el[0:10] for el in df["data"].values.tolist()]
    return df 

@app.route('/word')
def doc_template():
    doc = Document()
    doc.add_heading('Il mio primo report ', 0)
    folder = "static/report/"

    todisplaylist = ["andamento_nazionale.png","variazione_totale_positivi.png"]
    for name in todisplaylist:
        path = os.path.join(folder, name)
        doc.add_picture(path, width=Inches(6))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(file_stream, as_attachment=True, attachment_filename='report.docx',cache_timeout=0)

# https://iot-app-course.herokuapp.com/
@app.route('/')
def index():
    df = load_data()
    select = ["deceduti","totale_casi","dimessi_guariti"]
    fig1 = plot_plotly(df,x ="data", y=select,title="Andamento Nazionale")    
    fig2 = plot_plotly(df,x ="data", y=["variazione_totale_positivi"],title="Variazione totale positivi")   
    return render_template('index.html', 
                            name="Manuel",
                            surname="Rucci",
                            eta="26",
                            sesso="Maschio",
							collaboratore = "Sono Ugo e collaboro all'app",
                            fig1= convert_plotly_fig_to_json(fig1),
                            fig2= convert_plotly_fig_to_json(fig2),
                            )

if __name__ == "__main__":
    app.run()
